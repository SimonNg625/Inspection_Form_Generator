import streamlit as st
import os
import random
import calendar
import io
import zipfile
from datetime import date, timedelta
from docx import Document

# ==========================================
# 1. HONG KONG HOLIDAYS DATABASE
# ==========================================
class HKHolidays:
    """Contains public holiday dates for Hong Kong (2024-2026)."""
    HOLIDAYS = {
        # --- 2024 ---
        date(2024, 1, 1), date(2024, 2, 10), date(2024, 2, 12), date(2024, 2, 13),
        date(2024, 3, 29), date(2024, 3, 30), date(2024, 4, 1), date(2024, 4, 4),
        date(2024, 5, 1), date(2024, 5, 15), date(2024, 6, 10), date(2024, 7, 1),
        date(2024, 9, 18), date(2024, 10, 1), date(2024, 10, 11), date(2024, 12, 25), date(2024, 12, 26),
        # --- 2025 ---
        date(2025, 1, 1), date(2025, 1, 29), date(2025, 1, 30), date(2025, 1, 31),
        date(2025, 4, 4), date(2025, 4, 18), date(2025, 4, 19), date(2025, 4, 21),
        date(2025, 5, 1), date(2025, 5, 5), date(2025, 5, 31), date(2025, 7, 1),
        date(2025, 10, 1), date(2025, 10, 7), date(2025, 10, 29), date(2025, 12, 25), date(2025, 12, 26),
        # --- 2026 (Projected) ---
        date(2026, 1, 1), date(2026, 2, 17), date(2026, 2, 18), date(2026, 2, 19),
        date(2026, 4, 3), date(2026, 4, 4), date(2026, 4, 6), date(2026, 5, 1),
        date(2026, 5, 25), date(2026, 6, 19), date(2026, 7, 1), date(2026, 9, 26),
        date(2026, 10, 1), date(2026, 10, 19), date(2026, 12, 25), date(2026, 12, 26),
    }

    @staticmethod
    def is_holiday(check_date):
        return check_date in HKHolidays.HOLIDAYS

# ==========================================
# 2. CONFIGURATION & UTILS
# ==========================================
class Config:
    VALID_INSP_TYPES = ["RGI", "SAFE", "WSIN", "ENVI"]
    KEY_MAPPING = {
        "PI Inspection Form": "location", "Project No": "project_no",
        "Inspector": "inspector", "Contractor": "contractor",
        "Form No": "form_no", "Insp Type": "insp_type",
        "Scheduled": "scheduled", "Deadline": "deadline",
        "Performed By": "inspector", "Checked By": "checker",
        "Date": "perform_date",
    }

class DocUtils:
    @staticmethod
    def find_next_real_cell(row_cells, current_index):
        current_cell = row_cells[current_index]
        next_index = current_index + 1
        while next_index < len(row_cells):
            next_cell = row_cells[next_index]
            if next_cell._element is not current_cell._element:
                return next_cell
            next_index += 1
        return None

    @staticmethod
    def safe_update_cell(cell, new_value):
        if not cell: return
        if cell.paragraphs:
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].text = str(new_value)
                for run in p.runs[1:]: run.text = ""
            else:
                p.add_run(str(new_value))
        else:
            cell.text = str(new_value)

class DateEngine:
    @staticmethod
    def get_random_weekday(year, month, start_day, end_day):
        _, last_day_of_month = calendar.monthrange(year, month)
        actual_start = max(1, start_day)
        actual_end = min(end_day, last_day_of_month)
        
        if actual_start > actual_end: return date(year, month, actual_end)

        valid_dates = []
        for day in range(actual_start, actual_end + 1):
            try:
                curr_date = date(year, month, day)
                # Weekday < 5 means Mon(0) to Fri(4)
                if curr_date.weekday() < 5 and not HKHolidays.is_holiday(curr_date):
                    valid_dates.append(curr_date)
            except ValueError: continue
            
        if valid_dates: return random.choice(valid_dates)
        return date(year, month, actual_start)

# ==========================================
# 3. EXTRACTION LOGIC
# ==========================================
def extract_details_from_doc(doc_obj):
    extracted = {
        "location": "", "project_no": "",
        "inspector": "", "contractor": "", "checker": ""
    }
    keys = {
        "PI Inspection Form": "location", "Location": "location",
        "Project No": "project_no", "Inspector": "inspector",
        "Contractor": "contractor", "Checked By": "checker"
    }
    
    for table in doc_obj.tables:
        for row in table.rows:
            cells = row.cells
            for i, cell in enumerate(cells):
                cell_text = cell.text.strip()
                for key, field in keys.items():
                    if key in cell_text:
                        if key == "Inspector" and "PI Inspection" in cell_text: continue
                        target = DocUtils.find_next_real_cell(cells, i)
                        if target and target.text.strip():
                            extracted[field] = target.text.strip()
    return extracted

# ==========================================
# 4. GENERATION LOGIC
# ==========================================
def generate_docs_in_memory(template_file, form_data, start_date, end_date, req_type, start_num):
    # Load template into memory once to check validity
    try:
        template_bytes = template_file.getvalue()
    except:
        return None

    # Calculate Month Range
    start_m, start_y = start_date.month, start_date.year
    end_m, end_y = end_date.month, end_date.year
    
    generated_files = [] # List of (filename, bio_stream)
    counter = start_num
    
    curr_m, curr_y = start_m, start_y
    
    while (curr_y < end_y) or (curr_y == end_y and curr_m <= end_m):
        # Date Logic
        dates = []
        _, last_day = calendar.monthrange(curr_y, curr_m)
        
        if req_type == "SAFE":
            d1 = DateEngine.get_random_weekday(curr_y, curr_m, 1, 7)
            min_gap = d1 + timedelta(days=14)
            d2 = DateEngine.get_random_weekday(curr_y, curr_m, min_gap.day, last_day)
            dates = [d1, d2]
        else:
            dates = [DateEngine.get_random_weekday(curr_y, curr_m, 1, last_day)]
            
        # Prepare Base Data
        month_str = f"{curr_m:02d}"
        base_data = form_data.copy()
        base_data["scheduled"] = f"01/{month_str}/{curr_y}"
        base_data["deadline"] = f"{last_day}/{month_str}/{curr_y}"
        base_data["insp_type"] = req_type
        
        # Create Files
        for perform_dt in dates:
            date_str = perform_dt.strftime("%d/%m/%Y")
            form_no = f"IPRJ{req_type}{counter:04d}"
            
            final_data = base_data.copy()
            final_data["form_no"] = form_no
            final_data["perform_date"] = date_str
            
            # Open fresh doc from memory bytes
            doc = Document(io.BytesIO(template_bytes))
            
            # Fill Table
            for table in doc.tables:
                for row in table.rows:
                    cells = row.cells
                    for i, cell in enumerate(cells):
                        text = cell.text.strip()
                        for key, field in Config.KEY_MAPPING.items():
                            if key in text:
                                if key == "Inspector" and "PI Inspection" in text: continue
                                target = DocUtils.find_next_real_cell(cells, i)
                                if target:
                                    DocUtils.safe_update_cell(target, final_data.get(field, ""))
            
            # Save to memory stream
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            filename = f"{form_no}.docx"
            generated_files.append((filename, file_stream))
            
            counter += 1
            
        curr_m += 1
        if curr_m > 12:
            curr_m = 1
            curr_y += 1
            
    return generated_files

# ==========================================
# 5. STREAMLIT UI
# ==========================================
st.set_page_config(page_title="Insp Form Gen", page_icon="📝")

st.title("📝 Inspection Form Generator")
st.markdown("Upload a template, verify details, and generate a ZIP of filled forms.")

# --- STEP 1: UPLOAD ---
st.subheader("1. Upload Template")
uploaded_file = st.file_uploader("Upload Word Document (.docx)", type=["docx"])

if uploaded_file:
    # Auto-detect type from filename
    filename_upper = uploaded_file.name.upper()
    detected_type = "RGI" # Default
    for t in Config.VALID_INSP_TYPES:
        if t in filename_upper:
            detected_type = t
            break
            
    # Extract Data Preview
    if 'extracted_data' not in st.session_state:
        doc_preview = Document(uploaded_file)
        st.session_state['extracted_data'] = extract_details_from_doc(doc_preview)

    # --- STEP 2: CONFIGURATION ---
    st.divider()
    st.subheader("2. Configure Details")
    
    col1, col2 = st.columns(2)
    with col1:
        insp_type = st.selectbox("Inspection Type", Config.VALID_INSP_TYPES, index=Config.VALID_INSP_TYPES.index(detected_type) if detected_type in Config.VALID_INSP_TYPES else 0)
        start_num = st.number_input("Starting Form Number", min_value=1, value=1)
        
    with col2:
        # Date Range Picker
        today = date.today()
        d_range = st.date_input("Select Date Range (Start to End)", [today, today + timedelta(days=30)])
    
    st.info("Verify the data extracted from the template below:")
    
    # Editable Form for Extracted Data
    with st.expander("📂 Edit Project Details", expanded=True):
        e_data = st.session_state['extracted_data']
        
        c1, c2 = st.columns(2)
        new_loc = c1.text_input("Location", e_data.get("location"))
        new_proj = c2.text_input("Project No", e_data.get("project_no"))
        
        c3, c4 = st.columns(2)
        new_insp = c3.text_input("Inspector", e_data.get("inspector"))
        new_cont = c4.text_input("Contractor", e_data.get("contractor"))
        
        new_check = st.text_input("Checked By", e_data.get("checker"))
        
        # Update Dictionary
        form_data = {
            "location": new_loc, "project_no": new_proj,
            "inspector": new_insp, "contractor": new_cont,
            "checker": new_check
        }

    # --- STEP 3: GENERATE ---
    st.divider()
    if st.button("🚀 Generate Forms", type="primary"):
        if len(d_range) != 2:
            st.error("Please select both a Start Date and an End Date.")
        else:
            with st.spinner("Processing documents..."):
                start_d, end_d = d_range
                
                # Run Generation
                files_list = generate_docs_in_memory(
                    uploaded_file, form_data, start_d, end_d, insp_type, int(start_num)
                )
                
                if files_list:
                    # Create ZIP in memory
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, "w") as zf:
                        for fname, fcontent in files_list:
                            zf.writestr(fname, fcontent.getvalue())
                    
                    zip_buffer.seek(0)
                    
                    st.success(f"Success! Generated {len(files_list)} forms.")
                    
                    # Download Button
                    st.download_button(
                        label="📥 Download ZIP File",
                        data=zip_buffer,
                        file_name=f"Inspection_Forms_{insp_type}.zip",
                        mime="application/zip"
                    )
                else:
                    st.error("Failed to generate files.")
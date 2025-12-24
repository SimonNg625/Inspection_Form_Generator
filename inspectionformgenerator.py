import os
import random
import calendar
import shutil
import tkinter as tk
from tkinter import filedialog
from datetime import date, timedelta
from docx import Document

# --- CONFIGURATION ---

# 1. EXTENSIBILITY: Add new types here (e.g., "TEST")
VALID_INSP_TYPES = ["RGI", "SAFE", "WSIN", "ENVI"] 

KEY_MAPPING = {
    "PI Inspection Form": "location",
    "Project No": "project_no",
    "Inspector": "inspector",
    "Contractor": "contractor",
    "Form No": "form_no",
    "Insp Type": "insp_type",
    "Scheduled": "scheduled",
    "Deadline": "deadline",
    "Performed By": "inspector",   
    "Checked By": "checker",
    "Date": "perform_date",         
}

# --- HELPERS ---

def get_random_weekday_date_obj(year, month, start_day, end_day):
    """Returns random Mon-Fri date object between start and end day."""
    _, last_day_of_month = calendar.monthrange(year, month)
    actual_start = max(1, start_day)
    actual_end = min(end_day, last_day_of_month)
    
    if actual_start > actual_end:
        return date(year, month, actual_end)

    valid_dates = []
    for day in range(actual_start, actual_end + 1):
        try:
            curr_date = date(year, month, day)
            if curr_date.weekday() < 5: 
                valid_dates.append(curr_date)
        except ValueError:
            continue
            
    if valid_dates:
        return random.choice(valid_dates)
    return date(year, month, actual_start)

def find_next_real_cell(row_cells, current_index):
    """Skips merged cells to find the actual value cell."""
    current_cell = row_cells[current_index]
    next_index = current_index + 1
    while next_index < len(row_cells):
        next_cell = row_cells[next_index]
        if next_cell._element is not current_cell._element:
            return next_cell
        next_index += 1
    return None

def safe_update_cell(cell, new_value):
    """Updates cell text preserving format."""
    if not cell: return
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = str(new_value)
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.add_run(str(new_value))
    else:
        cell.text = str(new_value)

def extract_project_details(doc):
    """Reads the uploaded template to find existing Project/Location/Names."""
    extracted = {
        "location": "Unknown Location",
        "project_no": "Unknown Project",
        "inspector": "Unknown Inspector",
        "contractor": "Unknown Contractor",
        "checker": "Unknown Checker"
    }
    
    keys_to_extract = {
        "PI Inspection Form": "location", 
        "Location": "location",           
        "Project No": "project_no",
        "Inspector": "inspector",
        "Contractor": "contractor",
        "Checked By": "checker"
    }

    print("\n>> Scanning template for project details...")
    
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for i, cell in enumerate(cells):
                cell_text = cell.text.strip()
                for key, field_name in keys_to_extract.items():
                    if key in cell_text:
                        if key == "Inspector" and "PI Inspection" in cell_text:
                            continue
                        target_cell = find_next_real_cell(cells, i)
                        if target_cell:
                            val = target_cell.text.strip()
                            if val:
                                extracted[field_name] = val
                                print(f"   Found {field_name.upper()}: {val}")
    return extracted

def process_document(doc, data_dict):
    """Scans and updates tables."""
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for i, cell in enumerate(cells):
                cell_text = cell.text.strip()
                for key, data_field in KEY_MAPPING.items():
                    if key in cell_text:
                        if key == "Inspector" and "PI Inspection" in cell_text:
                            continue
                        target_cell = find_next_real_cell(cells, i)
                        if target_cell:
                            safe_update_cell(target_cell, data_dict[data_field])

def detect_type_from_filename(filename):
    """Checks filename against VALID_INSP_TYPES list."""
    fname_upper = os.path.basename(filename).upper()
    
    # Check against the list defined in CONFIGURATION
    for type_code in VALID_INSP_TYPES:
        if type_code in fname_upper:
            return type_code
            
    return "UNKNOWN"

def select_file_dialog():
    """Opens file explorer and returns path."""
    root = tk.Tk()
    root.withdraw() 
    root.attributes('-topmost', True) 
    filename = filedialog.askopenfilename(
        title="Select Inspection Form Template",
        filetypes=[("Word Documents", "*.docx")]
    )
    root.destroy()
    return filename

# --- MAIN EXECUTION ---

def main():
    print("--- Universal Inspection Form Generator (Strict Validation) ---")
    
    # --- NEW: OUTER LOOP TO ALLOW RESTARTING ---
    while True:
        
        # 1. INITIAL FILE SELECTION (With Loop for Validation)
        print("\nSTEP 1: Select your Template")
        
        template_filename = None
        template_type = "UNKNOWN"

        while True:
            template_filename = select_file_dialog()
            
            if not template_filename:
                print(">> No file selected. Exiting program.")
                return # Exits the entire function/program

            template_type = detect_type_from_filename(template_filename)

            # CHECK: Is the file type allowed?
            if template_type == "UNKNOWN":
                print(f"\n>> [!] ERROR: Invalid Filename.")
                print(f"   File '{os.path.basename(template_filename)}' does not contain a valid type keyword.")
                print(f"   Allowed types: {VALID_INSP_TYPES}")
                print("   Please rename your file or select the correct one.")
                input("   Press Enter to try again...") 
                continue 
            else:
                print(f">> Loaded '{os.path.basename(template_filename)}' ([{template_type}])")
                break 

        # 2. AUTO-EXTRACT DETAILS
        doc_for_scanning = Document(template_filename)
        project_data = extract_project_details(doc_for_scanning)
        
        print("\n>> Using the following details from the template:")
        print(f"   - Project: {project_data['project_no']}")
        print(f"   - Location: {project_data['location']}")
        
        # 3. TYPE SELECTION & VALIDATION LOOP
        requested_type = ""
        year = 2025

        while True:
            # Ask for input
            type_year_input = input(f"\nWhich Insp Type and year? (e.g., {template_type} 2025) [Press Enter for default]: ").strip()
            
            # --- STRICT INPUT PARSING ---
            if not type_year_input:
                requested_type = template_type
                year = 2025
                print(f">> Using default: {requested_type} {year}")
            else:
                try:
                    parts = type_year_input.split()
                    if len(parts) < 2:
                        raise ValueError("Missing Year")
                    
                    requested_type = parts[0].upper()
                    year = int(parts[1])
                    
                    if year < 1990 or year > 2100:
                        print(">> [!] Error: Please enter a realistic year (e.g., 2025).")
                        continue
                        
                except ValueError:
                    print(">> [!] Invalid format. Please enter 'TYPE YEAR' (e.g., SAFE 2025).")
                    continue 

            # --- CONFLICT CHECK ---
            if requested_type != template_type:
                print(f"\n[!] CONFLICT DETECTED")
                print(f"    You uploaded a [{template_type}] file.")
                print(f"    But you asked to generate [{requested_type}] forms.")
                
                # --- LOOP UNTIL VALID CHOICE IS MADE ---
                valid_fix = False
                while not valid_fix:
                    print("\nHow would you like to fix this?")
                    print("1. Re-upload the correct file")
                    print("2. Change my request")
                    
                    fix_choice = input("Enter 1 or 2: ").strip()
                    
                    if fix_choice == '1':
                        print("\n>> Opening file picker...")
                        
                        # Nested loop for re-upload validation
                        new_file_valid = False
                        while not new_file_valid:
                            new_file = select_file_dialog()
                            if not new_file:
                                print(">> No file selected. Keeping original file.")
                                new_file_valid = True # Exit re-upload loop, but didn't actually fix conflict
                            else:
                                new_type = detect_type_from_filename(new_file)
                                if new_type == "UNKNOWN":
                                    print(f">> [!] ERROR: Filename must contain: {VALID_INSP_TYPES}")
                                    input("Press Enter to try again...")
                                    continue
                                else:
                                    # Success
                                    template_filename = new_file
                                    template_type = new_type
                                    doc_for_scanning = Document(template_filename)
                                    project_data = extract_project_details(doc_for_scanning)
                                    print(f">> New file loaded: '{os.path.basename(template_filename)}' ([{template_type}])")
                                    new_file_valid = True
                        
                        valid_fix = True 

                    elif fix_choice == '2':
                        valid_fix = True 
                    
                    else:
                        print(">> [!] Invalid choice. Please enter '1' or '2' only.")
                
                continue # Restart Main Loop

            else:
                break 

        # 4. GENERATE & SAVE TO DOWNLOADS
        downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")
        output_dir_name = f"{requested_type}_{year}_Forms_Generated"
        output_dir = os.path.join(downloads_folder, output_dir_name)

        if os.path.exists(output_dir): 
            shutil.rmtree(output_dir)
        os.makedirs(output_dir)
        
        print(f"\n>> Generating {requested_type} forms in: {output_dir}")
        form_counter = 1

        for month in range(1, 13):
            month_str = f"{month:02d}"
            _, last_day_val = calendar.monthrange(year, month)
            scheduled = f"01/{month_str}/{year}"
            deadline = f"{last_day_val}/{month_str}/{year}"

            if requested_type == "SAFE":
                date1 = get_random_weekday_date_obj(year, month, 1, 7)
                min_gap_date = date1 + timedelta(days=14)
                date2 = get_random_weekday_date_obj(year, month, min_gap_date.day, last_day_val)
                dates_list = [date1, date2]
            else:
                dates_list = [get_random_weekday_date_obj(year, month, 1, last_day_val)]

            for perform_dt in dates_list:
                date_str = perform_dt.strftime("%d/%m/%Y")
                form_no = f"IPRJ{requested_type}{form_counter:04d}"
                
                final_data = project_data.copy()
                final_data.update({
                    "form_no": form_no,
                    "insp_type": requested_type,
                    "scheduled": scheduled,
                    "deadline": deadline,
                    "perform_date": date_str
                })

                doc = Document(template_filename)
                process_document(doc, final_data)
                doc.save(os.path.join(output_dir, f"{form_no}.docx"))
                print(f"Created: {form_no}.docx | {date_str}")
                form_counter += 1

        print(f"\nSUCCESS: Files saved to Downloads folder.")
        os.startfile(output_dir)
        
        # --- NEW: ASK TO RESTART OR EXIT ---
        print("\n" + "="*50)
        print("JOB COMPLETE. What would you like to do next?")
        print("1. Generate another type of inspection form")
        print("2. Exit Program")
        
        next_step = input("Enter 1 or 2: ").strip()
        
        if next_step == '1':
            print("\n" * 2) # Add some space
            print(">> Restarting program...")
            continue # Loops back to "STEP 1: Select your Template"
        else:
            print(">> Exiting. Goodbye!")
            break # Breaks the "while True" loop and ends program

if __name__ == "__main__":
    main()